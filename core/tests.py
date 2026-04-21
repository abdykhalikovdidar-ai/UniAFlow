from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required 
from django.contrib.auth import login
from django.db.models import Q, Sum, Count
from django.contrib.auth.models import User
from django.utils import timezone
from .models import Task, Profile, Message, Assignment
from .forms import RegisterForm
import json
import io
import calendar
from datetime import date, timedelta, datetime
from docx import Document

# --- 1. АВТОРИЗАЦИЯ И РЕГИСТРАЦИЯ ---

def register(request):
    if request.user.is_authenticated: 
        return redirect('index')
    if request.method == 'POST':
        form = RegisterForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('index')
    else:
        form = RegisterForm()
    return render(request, 'core/register.html', {'form': form})

# --- 2. ГЛАВНАЯ СТРАНИЦА (КАНБАН-ДОСКА) ---

@login_required
def index(request):
    profile_obj, _ = Profile.objects.get_or_create(user=request.user)
    profile_obj.last_seen = timezone.now()
    profile_obj.save()

    user_role = getattr(request.user.profile, 'role', 'student')
    query = request.GET.get('q', '')
    
    if user_role == 'teacher':
        course_filter = request.GET.get('course', '')
        group_filter = request.GET.get('group', '').strip()
        tasks = Task.objects.filter(status='done', is_verified=False).exclude(executor__is_superuser=True).order_by('-created_at')
        
        if query: 
            tasks = tasks.filter(Q(title__icontains=query) | Q(executor__username__icontains=query) | Q(executor__first_name__icontains=query) | Q(executor__last_name__icontains=query))
        if course_filter:
            tasks = tasks.filter(executor__profile__course=course_filter)
        if group_filter:
            tasks = tasks.filter(executor__profile__group_name__icontains=group_filter)

        return render(request, 'core/teacher_dashboard.html', {
            'tasks_to_review': tasks, 
            'query': query,
            'current_course': course_filter,
            'current_group': group_filter
        })
    else:
        tasks = Task.objects.filter(executor=request.user, is_verified=False)
        if query:
            tasks = tasks.filter(Q(title__icontains=query))
        unread = Message.objects.filter(receiver=request.user, is_read=False).count()
        return render(request, 'core/index.html', {'tasks': tasks, 'query': query, 'notifications_count': unread})

# --- 3. УПРАВЛЕНИЕ ПРОФИЛЕМ ---

@login_required
def profile(request):
    p_obj, _ = Profile.objects.get_or_create(user=request.user)
    user_role = p_obj.role
    
    if request.method == 'POST':
        if 'delete_image' in request.POST:
            p_obj.image = 'default.png'
            p_obj.save()
            return redirect('profile')

        request.user.first_name = request.POST.get('first_name', request.user.first_name)
        request.user.last_name = request.POST.get('last_name', request.user.last_name)
        new_email = request.POST.get('email')
        if new_email:
            request.user.email = new_email
        request.user.save()
        
        p_obj.phone_number = request.POST.get('phone_number', p_obj.phone_number)
        p_obj.bio = request.POST.get('bio', p_obj.bio)
        p_obj.gender = request.POST.get('gender', p_obj.gender)
        
        b_date = request.POST.get('birth_date')
        if b_date:
            p_obj.birth_date = b_date
        
        if user_role == 'teacher':
            p_obj.department = request.POST.get('department', p_obj.department)
            p_obj.academic_degree = request.POST.get('academic_degree', p_obj.academic_degree)
            p_obj.instruction_language = request.POST.get('instruction_language', p_obj.instruction_language)
            p_obj.experience = request.POST.get('experience', p_obj.experience)
            p_obj.office_number = request.POST.get('office_number', p_obj.office_number)
            p_obj.office_hours = request.POST.get('office_hours', p_obj.office_hours)
            p_obj.interests = request.POST.get('interests', p_obj.interests)
        else:
            p_obj.education_level = request.POST.get('education_level', p_obj.education_level)
            p_obj.course = request.POST.get('course', p_obj.course)
            p_obj.group_name = request.POST.get('group_name', p_obj.group_name)
            p_obj.language = request.POST.get('language', p_obj.language)

        if request.FILES.get('image'):
            p_obj.image = request.FILES.get('image')
        p_obj.save()
        return redirect('profile')

    template_name = 'core/teacher_profile.html' if user_role == 'teacher' else 'core/profile.html'
    return render(request, template_name, {'profile': p_obj})

# --- 4. ЗАДАНИЯ И ПРОЕКТЫ ---

@login_required
def assignments_page(request):
    user_role = getattr(request.user.profile, 'role', 'student')
    now = timezone.now()

    if user_role == 'teacher':
        if request.method == 'POST':
            deadline_val = request.POST.get('deadline')
            if deadline_val:
                deadline_val = datetime.strptime(deadline_val, '%Y-%m-%d %H:%M')
            
            Assignment.objects.create(
                title=request.POST.get('title'),
                description=request.POST.get('description'),
                teacher=request.user,
                target_course=request.POST.get('target_course', '').strip(),
                target_group=request.POST.get('target_group', '').strip(),
                difficulty=request.POST.get('difficulty', 'Средняя'),
                bonus_points=request.POST.get('bonus_points') or 0,
                deadline=deadline_val,
                max_grade=request.POST.get('max_grade') or 100,
                type=request.POST.get('type', 'word'),
                file=request.FILES.get('file')
            )
            return redirect('assignments_page')
        
        assignments = Assignment.objects.filter(teacher=request.user).annotate(taken_count=Count('tasks')).order_by('-created_at')
        return render(request, 'core/assignments.html', {'assignments': assignments, 'now': now})
    
    else:
        user_group = request.user.profile.group_name or ""
        user_course = request.user.profile.course or ""
        completed_assignment_ids = Task.objects.filter(executor=request.user, assignment__isnull=False, is_verified=True).values_list('assignment_id', flat=True)
        
        assignments = Assignment.objects.filter(
            (Q(target_group__iexact=user_group) | Q(target_group__exact='')) &
            (Q(target_course__iexact=user_course) | Q(target_course__exact='') | Q(target_course__icontains=user_course))
        ).exclude(id__in=completed_assignment_ids).order_by('-created_at')
        
        total_bonus = assignments.aggregate(Sum('bonus_points'))['bonus_points__sum'] or 0
        taken_ids = Task.objects.filter(executor=request.user, assignment__isnull=False).values_list('assignment_id', flat=True)
        
        return render(request, 'core/student_assignments.html', {
            'assignments': assignments, 
            'taken_ids': list(taken_ids),
            'total_bonus': total_bonus,
            'now': now
        })

@login_required
def take_assignment(request, assignment_id):
    assignment = get_object_or_404(Assignment, id=assignment_id)
    now = timezone.now()
    if request.user.profile.role == 'student':
        if assignment.deadline and assignment.deadline < now:
            return HttpResponse("Срок выполнения этого задания истек.", status=403)
        if not Task.objects.filter(executor=request.user, assignment=assignment).exists():
            Task.objects.create(
                title=assignment.title, 
                description=assignment.description, 
                executor=request.user, 
                assignment=assignment, 
                status='todo'
            )
    return redirect('index')

@login_required
def delete_assignment(request, assignment_id):
    assignment = get_object_or_404(Assignment, id=assignment_id)
    if request.user == assignment.teacher:
        assignment.delete()
    return redirect('assignments_page')

@login_required
def edit_assignment(request, assignment_id):
    assignment = get_object_or_404(Assignment, id=assignment_id)
    if request.user == assignment.teacher and request.method == 'POST':
        assignment.title = request.POST.get('title')
        assignment.description = request.POST.get('description')
        assignment.target_course = request.POST.get('target_course', '').strip()
        assignment.target_group = request.POST.get('target_group', '').strip()
        assignment.difficulty = request.POST.get('difficulty', assignment.difficulty)
        assignment.bonus_points = request.POST.get('bonus_points') or 0
        
        d_val = request.POST.get('deadline')
        if d_val:
            assignment.deadline = datetime.strptime(d_val, '%Y-%m-%d %H:%M')
            
        assignment.max_grade = request.POST.get('max_grade') or 100
        assignment.type = request.POST.get('type', 'word')
        if request.FILES.get('file'):
            assignment.file = request.FILES.get('file')
        assignment.save()
    return redirect('assignments_page')

@login_required
def copy_assignment(request, assignment_id):
    original = get_object_or_404(Assignment, id=assignment_id)
    if request.user == original.teacher:
        # Клонируем объект: создаем копию, сбрасываем ID и PK
        original.pk = None
        original.id = None
        original.title = f"{original.title} (Копия)"
        original.save()
    return redirect('assignments_page')

# --- 5. КАНБАН: УПРАВЛЕНИЕ ЗАДАЧАМИ ---

@login_required
def add_task(request):
    if request.method == "POST":
        title = request.POST.get('title')
        file = request.FILES.get('file')
        if title:
            Task.objects.create(title=title, executor=request.user, file=file, status='todo')
    return redirect('index')

@login_required
def edit_task(request, task_id):
    task = get_object_or_404(Task, id=task_id, executor=request.user)
    if request.method == "POST":
        task.title = request.POST.get('title')
        if request.FILES.get('file'):
            task.file = request.FILES.get('file')
        task.save()
        return redirect('index')
    return render(request, 'core/edit_task.html', {'task': task})

@login_required
def delete_task(request, task_id):
    task = get_object_or_404(Task, id=task_id, executor=request.user)
    task.delete()
    return redirect('index')

@csrf_exempt
@login_required
def update_task_status(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            task = Task.objects.get(id=data.get('id'), executor=request.user)
            task.status = data.get('status')
            if task.status == 'done':
                task.submitted_at = timezone.now()
            task.save()
            return JsonResponse({'success': True})
        except:
            return JsonResponse({'success': False})
    return JsonResponse({'success': False})

# --- 6. ПРОВЕРКА ЗАДАЧ (УЧИТЕЛЬ) ---

@login_required
def send_review(request, task_id):
    task = get_object_or_404(Task, id=task_id, executor=request.user)
    task.status = 'done'
    task.submitted_at = timezone.now()
    task.save()
    return redirect(request.META.get('HTTP_REFERER', 'index'))

@login_required
def return_task(request, task_id):
    task = get_object_or_404(Task, id=task_id, executor=request.user)
    if not task.is_verified:
        task.status = 'doing'
        task.save()
    return redirect(request.META.get('HTTP_REFERER', 'index'))

@login_required
def review_task(request, task_id):
    if request.user.profile.role == 'teacher' and request.method == 'POST':
        task = get_object_or_404(Task, id=task_id)
        action = request.POST.get('action')
        comment = request.POST.get('comment')
        grade_raw = request.POST.get('grade')
        
        try:
            grade = int(grade_raw) if grade_raw else 0
        except ValueError:
            grade = 0
            
        bonus = 0
        if task.assignment and task.submitted_at:
            if task.assignment.deadline and task.submitted_at <= task.assignment.deadline:
                bonus = int(task.assignment.bonus_points or 0)
        
        if action == 'approve': 
            task.is_verified = True
            task.grade = grade + bonus
            if bonus > 0:
                task.teacher_comment = f"{comment} (Бонус +{bonus} за скорость)" if comment else f"Бонус +{bonus} учтен."
            else:
                task.teacher_comment = comment or "Принято."
        elif action == 'reject': 
            task.status = 'doing'
            task.is_verified = False
            task.teacher_comment = comment or "Нужны правки ❌"
            
        task.save()
    return redirect('index')

# --- 7. ЖУРНАЛ И ОЦЕНКИ ---

@csrf_exempt
@login_required
def save_grade(request):
    if request.user.profile.role == 'teacher' and request.method == 'POST':
        try:
            data = json.loads(request.body)
            student_id = data.get('student_id')
            date_str = data.get('date')
            grade_val = data.get('grade')
            
            if grade_val == '':
                Task.objects.filter(executor_id=student_id, created_at__date=date_str, is_verified=True).delete()
                return JsonResponse({'status': 'ok', 'action': 'deleted'})
            
            new_grade = int(grade_val)
            task = Task.objects.filter(executor_id=student_id, created_at__date=date_str, is_verified=True).first()
            if not task:
                u = User.objects.get(id=student_id)
                task = Task.objects.create(title=f"Оценка за {date_str}", executor=u, status='done', is_verified=True, grade=new_grade)
                Task.objects.filter(id=task.id).update(created_at=date_str)
            else:
                task.grade = new_grade
                task.save()
            return JsonResponse({'status': 'ok', 'action': 'saved'})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)
    return JsonResponse({'status': 'error'}, status=400)

@login_required
def teacher_journal(request):
    if getattr(request.user.profile, 'role', 'student') != 'teacher':
        return redirect('index')
    
    course_filter = request.GET.get('course', '')
    group_filter = request.GET.get('group', '').strip()
    today = timezone.now().date()
    year = int(request.GET.get('year', today.year))
    month = int(request.GET.get('month', today.month))
    
    num_days = calendar.monthrange(year, month)[1]
    dates = [date(year, month, day) for day in range(1, num_days + 1) if date(year, month, day).weekday() < 5]
    
    students = User.objects.filter(profile__role='student', is_superuser=False)
    if course_filter: students = students.filter(profile__course=course_filter)
    if group_filter: students = students.filter(profile__group_name__icontains=group_filter)
    
    journal_data = []
    for student in students:
        total = Task.objects.filter(executor=student, is_verified=True).aggregate(Sum('grade'))['grade__sum'] or 0
        journal_data.append({'student': student, 'total_score': total})
        
    return render(request, 'core/teacher_journal.html', {
        'journal_data': journal_data, 
        'dates': dates, 
        'current_course': course_filter, 
        'current_group': group_filter,
        'year': year,
        'month': month
    })

# --- 8. СТАТИСТИКА И ЧАТ ---

@login_required
def stats(request):
    period = request.GET.get('period', 'month')
    all_tasks = Task.objects.filter(executor=request.user)
    now = timezone.now()
    
    if period == 'month':
        all_tasks = all_tasks.filter(created_at__gte=now - timedelta(days=30))
    elif period == 'autumn_2025':
        all_tasks = all_tasks.filter(created_at__range=["2025-09-01", "2025-11-30"])
    
    verified_tasks = all_tasks.filter(is_verified=True).exclude(grade__isnull=True)
    total_points = verified_tasks.aggregate(Sum('grade'))['grade__sum'] or 0
    
    return render(request, 'core/stats.html', {
        'todo_count': all_tasks.filter(status='todo').count(),
        'doing_count': all_tasks.filter(status='doing').count(),
        'done_count': all_tasks.filter(status='done').count(),
        'total_points': total_points,
        'grade_history': verified_tasks.order_by('-created_at'),
        'current_period': period
    })

@login_required
def messages_page(request):
    admin = User.objects.filter(is_superuser=True).first()
    user_role = getattr(request.user.profile, 'role', 'student')
    
    if request.method == 'POST' and admin:
        txt = request.POST.get('text', '').strip()
        if txt: Message.objects.create(sender=request.user, receiver=admin, text=txt)
        return redirect('messages_page')
        
    msgs = Message.objects.filter(Q(sender=request.user) | Q(receiver=request.user)).order_by('created_at')
    tasks = Task.objects.filter(executor=request.user, status='doing') if user_role == 'student' else []
    template = 'core/teacher_messages.html' if user_role == 'teacher' else 'core/messages.html'
    return render(request, template, {'msgs': msgs, 'admin_user': admin, 'tasks': tasks})

@login_required
def get_messages_api(request, user_id):
    target_user = get_object_or_404(User, id=user_id)
    Message.objects.filter(sender=target_user, receiver=request.user, is_read=False).update(is_read=True)
    
    msgs = Message.objects.filter(
        (Q(sender=request.user) & Q(receiver=target_user)) | 
        (Q(sender=target_user) & Q(receiver=request.user))
    ).order_by('created_at')
    
    data = [{
        'sender': m.sender.first_name or m.sender.username, 
        'text': m.text, 
        'is_my': m.sender == request.user, 
        'time': m.created_at.strftime("%H:%M")
    } for m in msgs]
    
    is_online = target_user.profile.is_online() if hasattr(target_user, 'profile') else False
    return JsonResponse({'messages': data, 'is_online': is_online})

@csrf_exempt
@login_required
def send_message_api(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        receiver = get_object_or_404(User, id=data.get('receiver_id'))
        msg = Message.objects.create(sender=request.user, receiver=receiver, text=data.get('text'))
        return JsonResponse({'status': 'ok', 'text': msg.text, 'time': msg.created_at.strftime("%H:%M")})
    return JsonResponse({'status': 'error'}, status=400)

@csrf_exempt
@login_required
def set_typing_status(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        p = request.user.profile
        p.is_typing = data.get('status', False)
        p.save()
        return JsonResponse({'status': 'ok'})
    return JsonResponse({'status': 'error'}, status=400)

# --- 9. ВОРД (ТРАНСКРИПТ) ---

@login_required
def download_transcript(request):
    tasks = Task.objects.filter(executor=request.user, is_verified=True).exclude(grade__isnull=True)
    total_points = tasks.aggregate(Sum('grade'))['grade__sum'] or 0
    profile = request.user.profile
    
    doc = Document()
    doc.add_heading('АКАДЕМИЧЕСКИЙ ТРАНСКРИПТ: UniFlow', 0)
    doc.add_paragraph(f'Студент: {request.user.get_full_name() or request.user.username}\nГруппа: {profile.group_name}\nДата: {date.today().strftime("%d.%m.%Y")}')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Дата'; hdr[1].text = 'Задание'; hdr[2].text = 'Балл'
    
    for t in tasks:
        row = table.add_row().cells
        row[0].text = t.created_at.strftime("%d.%m.%Y")
        row[1].text = t.title
        row[2].text = str(t.grade)
        
    doc.add_paragraph(f'\nИТОГО НАБРАНО: {total_points} баллов.')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    res = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    res['Content-Disposition'] = f'attachment; filename="Transcript_{request.user.username}.docx"'
    return res